import base64
import io
import json
import os
from tempfile import NamedTemporaryFile
from pdf2image import convert_from_path
from PIL import Image
from docx import Document
from pdf2image.exceptions import PDFPageCountError
from .file_handler import StreamlitFileHandler
from .prompt import PromptReader
import re
from time import sleep
from doc2pdf import convert
from concurrent.futures import ThreadPoolExecutor
from logger_config import logger
import subprocess
from PyPDF2 import PdfReader
from dotenv import load_dotenv
load_dotenv()
PATH_LIBREOFFICE = os.getenv("PATH_LIBREOFFICE") 

class ProcessDoc:
    def __init__(self, doc_file_uploaded, bedrock_client):
        self.doc_file_uploaded = doc_file_uploaded  # UploadedFile from Streamlit
        self.bedrock_client = bedrock_client
        self.model_id = "anthropic.claude-3-5-sonnet-20240620-v1:0"

    @staticmethod
    def prevent_table_split(doc_file_uploaded, output_path):
        """Ensure Word tables do not split across pages."""
        doc = Document(doc_file_uploaded)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    cell.paragraphs[0].paragraph_format.keep_together = True
                    cell.paragraphs[0].paragraph_format.keep_with_next = True
        doc.save(output_path)
        logger.info("executed prevent table split.")

    @staticmethod
    def convert_pdf_to_images(pdf_file):
        """Convert a PDF file into individual images."""
        with NamedTemporaryFile(delete=False, suffix=".pdf") as temp_pdf:
            temp_pdf.write(pdf_file.read())
            temp_pdf_path = temp_pdf.name
            logger.info(f"uploaded file into temp pdf file: {temp_pdf_path}")
        try:
            # Convert the PDF to images
            images = convert_from_path(temp_pdf_path, dpi=300)
            logger.info("converte temp pdf to images.")
        except PDFPageCountError:
            raise ValueError("The uploaded file is not a valid PDF or is corrupted.")
        except Exception as e:
            logger.info(f"Error occured: {e}")
        finally:
            # Clean up the temporary file
            os.remove(temp_pdf_path)
            logger.info("removed temp pdf file.")

        return images  # Return only the first 10 images
    
    @staticmethod
    def convert_to_pdf(input_file, output_folder):
        command = [
            "libreoffice",
            "--headless",
            "--convert-to", "pdf",
            "--outdir", output_folder,
            input_file,
        ]
        result = subprocess.run(command, stdout=subprocess.PIPE, stderr=subprocess.PIPE)
        logger.info(f"converted to pdf: {output_folder}")
        if result.returncode != 0:
            raise RuntimeError(f"LibreOffice conversion failed: {result.stderr.decode('utf-8')}")
        return f"File converted successfully to {output_folder}"
        
    @staticmethod
    def convert_docx_to_images(docx_file):
        """Convert a DOCX file into individual images via PDF."""
        with NamedTemporaryFile(delete=False, suffix=".docx") as temp_docx:
            temp_docx.write(docx_file.read())
            temp_docx_path = temp_docx.name
            logger.info(f"saved uploaded docx in temp -: {temp_docx_path}")

        try:
            ProcessDoc.prevent_table_split(temp_docx_path, temp_docx_path)
            pdf_path = temp_docx_path.replace(".docx", ".pdf")
            ProcessDoc.convert_to_pdf(temp_docx_path, os.path.dirname(pdf_path))
            if not os.path.isfile(pdf_path) or os.path.getsize(pdf_path) == 0:
                raise ValueError("DOCX to PDF conversion failed or produced an invalid PDF.")

            try:
                # Validate PDF
                reader = PdfReader(pdf_path)
                if not reader.pages:
                    raise ValueError("PDF has no pages.")

                # Convert PDF to images
                images = convert_from_path(pdf_path, dpi=300)
                logger.info("converted pdf to image.")
            except Exception as e:
                logger.error("Error occured: {e}")
                raise ValueError(f"PDF validation or image conversion failed: {e}")
        except Exception as e:
            logger.error(e)
        # finally:
        #     # Cleanup
        #     for file_path in [temp_docx_path, pdf_path]:
        #         if os.path.isfile(file_path):
        #             try:
        #                 os.remove(file_path)
        #                 logger.info("removed all temp files.")
        #             except Exception as e:
        #                 logger.error(f"Warning: Failed to remove temp file {file_path}: {e}")
        return images


    @staticmethod
    def image_to_base64(image):
        """Convert an image to base64 encoding."""
        with Image.new("RGB", image.size, "white") as bg:
            try:
                bg.paste(image)
                buffer = io.BytesIO()
                bg.save(buffer, format="JPEG")
                buffer.seek(0)
                image_b64_string = base64.b64encode(buffer.read()).decode("utf-8")
                logger.info("converted image to base64.")
                return image_b64_string
            except Exception as e:
                logger.error("Error occured converting image to base64: {e}")
                raise e

    def extract_number(self,value):
        try:
            match = re.search(r'\d+', value)  # Extract the first sequence of digits

            return match.group(0) if match else None
        except Exception as e:
            logger.error("Error occured in extracting number: {e}")
            raise e 
        
    
    def process_files_data(self):
        # Determine file type and convert to images
        file_name = self.doc_file_uploaded.name.lower()
        print("uploaded file name----: ",file_name)
        logger.info(f"uploaded file name----: {file_name}")
        if file_name.endswith('.pdf'):
            images = self.convert_pdf_to_images(self.doc_file_uploaded)
        elif file_name.endswith('.docx'):
            images = self.convert_docx_to_images(self.doc_file_uploaded)
        else:
            raise ValueError("Unsupported file type. Please upload a PDF or DOCX file.")
        return images

    def process_il_cr_numbers(self,images):
        """Process the uploaded files and get response from the model."""
        # Read the prompt
        pr = PromptReader('prompts/cr_il_from_loan_application_prompt.txt')
        prompt = pr.read_prompt()
        # Convert the first image to base64
        image_b64 = self.image_to_base64(images[2])

        # Get the response from the model
        doc_response = self.bedrock_client.get_response(
            prompt=prompt,
            encoded_file=image_b64,
            mime_type="image/jpeg",
            model_id=self.model_id,
        )
        doc_response = json.loads(doc_response)
        logger.info(f'type of doc_response: {type(doc_response)}')
        logger.info(f"response fro bedrock for cr and IL number: {doc_response}")
        commercial_register_value = None
        industrial_license_value = None
        for key,value in doc_response.items():
            if "commercial" in key.lower():
                commercial_register_value  = self.extract_number(value)
            if "industr" in key.lower():
                industrial_license_value = self.extract_number(value)
          
        logger.info(f"CR,IL from loan application. commercial_register_value:{commercial_register_value}, industrial_license_value:{industrial_license_value}")
        return {"commercial_register_value":commercial_register_value, "industrial_license_value":industrial_license_value}
    
    def missing_fileds(self,images):
        """Process the uploaded files and get response from the model."""
        result = []
        pr = PromptReader('prompts/missing.txt')
        prompt_missing_data = pr.read_prompt()
        logger.info("read propmt from prompts/missing.txt")
        # Convert the first image to base64
        for image in images[-4:]:
            image_b64 = self.image_to_base64(image)

        # Get the response from the model
            fields_data = self.bedrock_client.get_response(
                prompt=prompt_missing_data,
                encoded_file=image_b64,
                mime_type="image/jpeg",
                model_id=self.model_id,
            )
            result.append(fields_data)
        logger.info(f"result from sonnet for missing data: {result}")
        
        pr = PromptReader('prompts/missing_result.txt')
        prompt_missing_result = pr.read_prompt()
        logger.info("read prompt from prompts/missing_result.txt")
        missing_data = self.bedrock_client.get_response_text(prompt_missing_result,result,model_id=self.model_id)
        logger.info(f"response from bedrock for missing fileds: {missing_data}")        
        return {"file_data":result,
                "missing_data":missing_data
                }

    def process_inconsistent_data(self,images):
        result = []
        pr = PromptReader('prompts/loan_application.txt')
        prompt_full_data = pr.read_prompt()
        logger.info("read propmt from prompts/loan_application.txt")
        # Convert the first image to base64
        for image in images:
            image_b64 = self.image_to_base64(image)

        # Get the response from the model
            fields_data = self.bedrock_client.get_response(
                prompt=prompt_full_data,
                encoded_file=image_b64,
                mime_type="image/jpeg",
                model_id=self.model_id,
            )
            result.append(fields_data)
            logger.info("extracted data from image.")
        pr = PromptReader('prompts/inconsistent_prompt.txt')
        prompt_inconsistent_data = pr.read_prompt()
        logger.info("read propmt from prompts/inconsistent_prompt.txt")
        
        # Get the response from the model
        incs_data = self.bedrock_client.get_response_text(prompt_inconsistent_data,result,model_id=self.model_id)
        logger.info(f"result from sonnet for inconsistent data: {fields_data}")
        return incs_data

    def file_processor(self):
        images = ProcessDoc.process_files_data(self)
        # Define the parallel tasks
        with ThreadPoolExecutor() as executor:
            future_resp_1 = executor.submit(self.process_il_cr_numbers, images)
            future_resp_2 = executor.submit(self.missing_fileds, images)
            future_resp_3 = executor.submit(self.process_inconsistent_data,images)
            # Retrieve results from futures
            resp_1 = future_resp_1.result()
            resp_2 = future_resp_2.result()
            resp_3 = future_resp_3.result()
        final_result = {"cr_and_il": resp_1, "missing": resp_2,"inc_data":resp_3}
        logger.info(f"IL,CR,missing data and inconsistent data: {final_result}")
        return final_result

