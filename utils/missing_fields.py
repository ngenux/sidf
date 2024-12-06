import base64
import io
import json
import os
from tempfile import NamedTemporaryFile
from pdf2image import convert_from_path
from PIL import Image
from docx import Document
from pdf2image.exceptions import PDFPageCountError
from .file_handler import StreamlitFileHandler
from .prompt import PromptReader
import boto3
import re
from time import sleep
from pdf2image.exceptions import PDFPageCountError
from doc2pdf import convert

class ProcessMissingfields:
    def __init__(self, doc_file_uploaded, bedrock_client):
        self.doc_file_uploaded = doc_file_uploaded  # UploadedFile from Streamlit
        self.bedrock_client = bedrock_client
        self.model_id = "anthropic.claude-3-5-sonnet-20240620-v1:0"

    @staticmethod
    def prevent_table_split(doc_file_uploaded, output_path):
        """Ensure Word tables do not split across pages."""
        doc = Document(doc_file_uploaded)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    cell.paragraphs[0].paragraph_format.keep_together = True
                    cell.paragraphs[0].paragraph_format.keep_with_next = True
        doc.save(output_path)

    @staticmethod
    def convert_pdf_to_images(pdf_file):
        """Convert a PDF file into individual images."""
        with NamedTemporaryFile(delete=False, suffix=".pdf") as temp_pdf:
            temp_pdf.write(pdf_file.read())
            temp_pdf_path = temp_pdf.name

        try:
            # Convert the PDF to images
            images = convert_from_path(temp_pdf_path, dpi=300)
            print('converted pdf to images')
        except PDFPageCountError:
            raise ValueError("The uploaded file is not a valid PDF or is corrupted.")
        finally:
            # Clean up the temporary file
            os.remove(temp_pdf_path)

        return images 

    @staticmethod
    def convert_docx_to_images(docx_file):
        """Convert a DOCX file into individual images via PDF."""
        with NamedTemporaryFile(delete=False, suffix=".docx") as temp_docx:
            try:
                # Save the uploaded DOCX file to a temporary location
                temp_docx.write(docx_file.read())
                temp_docx_path = temp_docx.name
                print("temp doc file name---: ",temp_docx_path)
                # Prevent table splitting and save to the same temporary DOCX file
                ProcessMissingfields.prevent_table_split(temp_docx_path, temp_docx_path)
                print("processed table split")
                # Convert the processed DOCX to PDF
                pdf_path = temp_docx_path.replace(".docx",".pdf")
                convert(temp_docx_path, pdf_path)
                print('converted temp doc to pdf')
                # Validate the generated PDF
                if not os.path.exists(pdf_path) or os.path.getsize(pdf_path) == 0:
                    raise ValueError("Failed to convert DOCX to PDF. The file is empty or invalid.")

                # Convert the PDF to images
                images = convert_from_path(pdf_path, dpi=300)
                print('converted pdf to images')
            except PDFPageCountError:
                raise ValueError("Unable to process the PDF file generated from DOCX.")
            except Exception as e:
                raise ValueError(f"Error converting DOCX to images: {e}")
            finally:
                # Ensure temporary files are removed
                for file_path in [temp_docx_path, pdf_path]:
                    for _ in range(3):  # Retry mechanism
                        try:
                            os.remove(file_path)
                            break
                        except PermissionError:
                            sleep(0.1)  # Delay before retrying
                        except Exception as e:
                            print(f"Warning: Failed to remove temp file {file_path}: {e}")

        return images


    @staticmethod
    def image_to_base64(image):
        """Convert an image to base64 encoding."""
        with Image.new("RGB", image.size, "white") as bg:
            bg.paste(image)
            buffer = io.BytesIO()
            bg.save(buffer, format="JPEG")
            buffer.seek(0)
            image_b64_string = base64.b64encode(buffer.read()).decode("utf-8")
            return image_b64_string

    @staticmethod
    def get_messages(image_b64_string, prompt):
        """Prepare the messages payload."""
        messages = [
            {
                "role": "user",
                "content": [
                    {"type": "text", "text": "<image>"},
                    {
                        "type": "image",
                        "source": {
                            "type": "base64",
                            "media_type": "image/jpeg",
                            "data": image_b64_string,
                        },
                    },
                    {"type": "text", "text": "</image>"},
                    {"type": "text", "text": prompt},
                ],
            }
        ]
        return messages

    @staticmethod
    def get_body(messages):
        """Create the request body for the model."""
        body = json.dumps(
            {
                "anthropic_version": "bedrock-2023-05-31",
                "max_tokens": 4000,
                "temperature": 0.1,
                "top_k": 250,
                "top_p": 0.999,
                "stop_sequences": ["\n\nHuman"],
                "messages": messages,
            }
        )
        return body


    def process_files(self):
        """Process the uploaded files and get response from the model."""
        # Read the prompt
        pr = PromptReader('prompts/loan_application.txt')
        prompt = pr.read_prompt()

        # Determine file type and convert to images
        file_name = self.doc_file_uploaded.name.lower()
        print("uploaded file name----: ",file_name)
        if file_name.endswith('.pdf'):
            images = self.convert_pdf_to_images(self.doc_file_uploaded)
        elif file_name.endswith('.docx'):
            images = self.convert_docx_to_images(self.doc_file_uploaded)
        else:
            raise ValueError("Unsupported file type. Please upload a PDF or DOCX file.")

        # Convert the first image to base64
        result = []
        for image in images:
            image_b64 = self.image_to_base64(image)
            # Get the response from the model
            doc_response = self.bedrock_client.get_response(
                prompt=prompt,
                encoded_file=image_b64,
                mime_type="image/jpeg",
                model_id=self.model_id,
            )
            result.append(doc_response)
        print(result)
        return result

