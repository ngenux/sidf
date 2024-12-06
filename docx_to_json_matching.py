import os
import base64
import json
import boto3
import io
from PIL import Image
from pdf2image import convert_from_path
from docx import Document
from docx2pdf import convert
from logger_config import logger
from prompt import prompt
from dotenv import load_dotenv
import os

load_dotenv()
# Initialize boto3 client for Bedrock
# aws_secret_access_key="yonTteFP85W4DcGQkGY0IHVb1WwrNsfA6jqqoITM" ,
#                 aws_access_key_id="AKIARSIADPBUXXXFF4MY",
bedrock = boto3.client(service_name="bedrock-runtime",aws_access_key_id=os.getenv("AWS_ACCESS_KEY_ID"),
                            aws_secret_access_key=os.getenv("AWS_SECRET_ACCESS_KEY"), region_name="us-east-1")

# Prevent table splits across pages in the DOCX file
def prevent_table_split(doc_path, output_path):
    doc = Document(doc_path)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                cell.paragraphs[0].paragraph_format.keep_together = True
                cell.paragraphs[0].paragraph_format.keep_with_next = True
    doc.save(output_path)


# Convert a PDF file into individual images
def convert_pdf_to_images(pdf_path):
    images = convert_from_path(pdf_path, dpi=300)
    return images  # Return only the first 10 images


# Convert an image to base64 encoding
def image_to_base64(image):
    with Image.new("RGB", image.size, "white") as bg:
        bg.paste(image)
        buffer = io.BytesIO()
        bg.save(buffer, format="JPEG")
        buffer.seek(0)
        image_b64_string = base64.b64encode(buffer.read()).decode("utf-8")
        return image_b64_string


# Prepare the messages payload
def get_messages(image_b64_string):
    messages = [
        {
            "role": "user",
            "content": [
                {"type": "text", "text": "<image>"},
                {
                    "type": "image",
                    "source": {
                        "type": "base64",
                        "media_type": "image/jpeg",
                        "data": image_b64_string,
                    },
                },
                {"type": "text", "text": "</image>"},
                {"type": "text", "text": prompt},
            ],
        }
    ]
    return messages


# Create the request body
def get_body(messages):
    body = json.dumps(
        {
            "anthropic_version": "bedrock-2023-05-31",
            "max_tokens": 4000,
            "temperature": 0.1,
            "top_k": 250,
            "top_p": 0.999,
            "stop_sequences": ["\n\nHuman"],
            "messages": messages,
        }
    )
    return body


# Fetch the response from the Bedrock model
def get_response(body):
    # modelId = "anthropic.claude-3-haiku-20240307-v1:0"

    model_id = "anthropic.claude-3-5-sonnet-20240620-v1:0"

    content_type = "application/json"
    accept = "application/json"

    response = bedrock.invoke_model(
        modelId=model_id, contentType=content_type, accept=accept, body=body
    )
    response_body = json.loads(response.get("body").read())
    result_dict = {
        "text": response_body["content"][0]["text"],
        "tokens": response_body["usage"],
    }
    return result_dict


# Main function
def main_func(args):
    input_path = args["input_path"]
    file_extension = os.path.splitext(input_path)[1].lower()

    if file_extension == ".docx":
        logger.info("Processing DOCX file...")
        # Step 1: Prevent table split and save updated DOCX
        updated_docx_path = "updated_" + os.path.basename(input_path)
        prevent_table_split(input_path, updated_docx_path)
        logger.info("Successfully prevented the table split")

        # Step 2: Convert updated DOCX to PDF
        pdf_path = updated_docx_path.replace(".docx", ".pdf")
        convert(updated_docx_path, pdf_path)
        logger.info("Successfully converted the DOCX to a PDF file")
    elif file_extension == ".pdf":
        logger.info("Processing PDF file...")
        pdf_path = input_path
    else:
        raise ValueError("Unsupported file type. Please provide a DOCX or PDF file.")

    # Step 3: Convert PDF to the first 10 images
    images = convert_pdf_to_images(pdf_path)
    logger.info("Converted PDF to images")

    # Step 4: Convert images to base64 and get responses
    results = []
    for image in images:
        image_b64 = image_to_base64(image)
        messages = get_messages(image_b64)
        body = get_body(messages)
        response = get_response(body)
        print(response)
        results.append(response)

    logger.info("Processed all images and fetched responses")
    return {"file_path": input_path, "results": results}


# Example usage
if __name__ == "__main__":
    args = {
        "input_path": "Industrial Loan Application Form (Ceramica Terra).docx"
    }  # Change to .pdf to test PDF flow
    responses = main_func(args)
    # print(responses['content'])
    data = responses["results"]
    print(data)
    res = json.loads(data[2]["text"])
    commercial_register = res["GENERAL INFORMATION"][
        "Commercial Register No. and date"
    ].split(" / ")[0]
    industrial_license = res["GENERAL INFORMATION"][
        "Industrial license no. and date"
    ].split(" / ")[0]

    print(f"Commercial Registration Number: {commercial_register}")
    print(f"Industrial License Number: {industrial_license}")